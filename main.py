from datetime import datetime
import os
import re
import shutil
from urllib.request import urlopen, Request
import requests
from bs4 import BeautifulSoup
from docx import Document
from pdf2docx import Converter
from py7zr import unpack_7zarchive
from pyunpack import Archive
import pandas as pd
from win32com import client as wc
from PySide6.QtCore import (QCoreApplication, QDate, QDateTime, QLocale,
    QMetaObject, QObject, QPoint, QRect,
    QSize, QTime, QUrl, Qt)
from PySide6.QtGui import (QBrush, QColor, QConicalGradient, QCursor,
    QFont, QFontDatabase, QGradient, QIcon,
    QImage, QKeySequence, QLinearGradient, QPainter,
    QPalette, QPixmap, QRadialGradient, QTransform)
from PySide6.QtWidgets import (QApplication, QComboBox, QMainWindow, QMenuBar,
    QPushButton, QSizePolicy, QStatusBar, QWidget)

class Ui_MainWindow(object):
    def setupUi(self, MainWindow):
        if not MainWindow.objectName():
            MainWindow.setObjectName(u"MainWindow")
        MainWindow.resize(800, 600)
        self.centralwidget = QWidget(MainWindow)
        self.centralwidget.setObjectName(u"centralwidget")
        self.pushButton = QPushButton(self.centralwidget)
        self.pushButton.setObjectName(u"pushButton")
        self.pushButton.setGeometry(QRect(620, 70, 75, 24))
        self.comboBox = QComboBox(self.centralwidget)
        self.comboBox.addItem("")
        self.comboBox.addItem("")
        self.comboBox.addItem("")
        self.comboBox.addItem("")
        self.comboBox.addItem("")
        self.comboBox.addItem("")
        self.comboBox.addItem("")
        self.comboBox.addItem("")
        self.comboBox.addItem("")
        self.comboBox.addItem("")
        self.comboBox.addItem("")
        self.comboBox.addItem("")
        self.comboBox.addItem("")
        self.comboBox.addItem("")
        self.comboBox.addItem("")
        self.comboBox.addItem("")
        self.comboBox.addItem("")
        self.comboBox.addItem("")
        self.comboBox.addItem("")
        self.comboBox.addItem("")
        self.comboBox.addItem("")
        self.comboBox.setObjectName(u"comboBox")
        self.comboBox.setGeometry(QRect(20, 70, 441, 22))
        MainWindow.setCentralWidget(self.centralwidget)
        self.menubar = QMenuBar(MainWindow)
        self.menubar.setObjectName(u"menubar")
        self.menubar.setGeometry(QRect(0, 0, 800, 22))
        MainWindow.setMenuBar(self.menubar)
        self.statusbar = QStatusBar(MainWindow)
        self.statusbar.setObjectName(u"statusbar")
        MainWindow.setStatusBar(self.statusbar)

        self.retranslateUi(MainWindow)

        QMetaObject.connectSlotsByName(MainWindow)
    # setupUi

    def retranslateUi(self, MainWindow):
        MainWindow.setWindowTitle(QCoreApplication.translate("MainWindow", u"MainWindow", None))
        self.pushButton.setText(QCoreApplication.translate("MainWindow", u"\u041d\u0430\u0439\u0442\u0438", None))
        self.comboBox.setItemText(0, QCoreApplication.translate("MainWindow", u"A. \u041f\u0440\u043e\u0434\u0443\u043a\u0446\u0438\u044f \u0441\u0435\u043b\u044c\u0441\u043a\u043e\u0433\u043e, \u043b\u0435\u0441\u043d\u043e\u0433\u043e \u0438 \u0440\u044b\u0431\u043d\u043e\u0433\u043e \u0445\u043e\u0437\u044f\u0439\u0441\u0442\u0432\u0430", None))
        self.comboBox.setItemText(1, QCoreApplication.translate("MainWindow", u"B. \u041f\u0440\u043e\u0434\u0443\u043a\u0446\u0438\u044f \u0433\u043e\u0440\u043d\u043e\u0434\u043e\u0431\u044b\u0432\u0430\u044e\u0449\u0438\u0445 \u043f\u0440\u043e\u0438\u0437\u0432\u043e\u0434\u0441\u0442\u0432", None))
        self.comboBox.setItemText(2, QCoreApplication.translate("MainWindow", u"C. \u041f\u0440\u043e\u0434\u0443\u043a\u0446\u0438\u044f \u043e\u0431\u0440\u0430\u0431\u0430\u0442\u044b\u0432\u0430\u044e\u0449\u0438\u0445 \u043f\u0440\u043e\u0438\u0437\u0432\u043e\u0434\u0441\u0442\u0432", None))
        self.comboBox.setItemText(3, QCoreApplication.translate("MainWindow", u"D. \u042d\u043b\u0435\u043a\u0442\u0440\u043e\u044d\u043d\u0435\u0440\u0433\u0438\u044f, \u0433\u0430\u0437, \u043f\u0430\u0440 \u0438 \u043a\u043e\u043d\u0434\u0438\u0446\u0438\u043e\u043d\u0438\u0440\u043e\u0432\u0430\u043d\u0438\u0435 \u0432\u043e\u0437\u0434\u0443\u0445\u0430", None))
        self.comboBox.setItemText(4, QCoreApplication.translate("MainWindow", u"E. \u0412\u043e\u0434\u043e\u0441\u043d\u0430\u0431\u0436\u0435\u043d\u0438\u0435; \u0432\u043e\u0434\u043e\u043e\u0442\u0432\u0435\u0434\u0435\u043d\u0438\u0435, \u0443\u0441\u043b\u0443\u0433\u0438 \u043f\u043e \u0443\u0434\u0430\u043b\u0435\u043d\u0438\u044e \u0438 \u0440\u0435\u043a\u0443\u043b\u044c\u0442\u0438\u0432\u0430\u0446\u0438\u0438 \u043e\u0442\u0445\u043e\u0434\u043e\u0432", None))
        self.comboBox.setItemText(5, QCoreApplication.translate("MainWindow", u"F. \u0421\u043e\u043e\u0440\u0443\u0436\u0435\u043d\u0438\u044f \u0438 \u0441\u0442\u0440\u043e\u0438\u0442\u0435\u043b\u044c\u043d\u044b\u0435 \u0440\u0430\u0431\u043e\u0442\u044b", None))
        self.comboBox.setItemText(6, QCoreApplication.translate("MainWindow", u"G. \u0423\u0441\u043b\u0443\u0433\u0438 \u043f\u043e \u043e\u043f\u0442\u043e\u0432\u043e\u0439 \u0438 \u0440\u043e\u0437\u043d\u0438\u0447\u043d\u043e\u0439 \u0442\u043e\u0440\u0433\u043e\u0432\u043b\u0435; \u0443\u0441\u043b\u0443\u0433\u0438 \u043f\u043e \u0440\u0435\u043c\u043e\u043d\u0442\u0443 \u0430\u0432\u0442\u043e\u0442\u0440\u0430\u043d\u0441\u043f\u043e\u0440\u0442\u043d\u044b\u0445 \u0441\u0440\u0435\u0434\u0441\u0442\u0432 \u0438 \u043c\u043e\u0442\u043e\u0446\u0438\u043a\u043b\u043e\u0432", None))
        self.comboBox.setItemText(7, QCoreApplication.translate("MainWindow", u"H. \u0423\u0441\u043b\u0443\u0433\u0438 \u0442\u0440\u0430\u043d\u0441\u043f\u043e\u0440\u0442\u0430 \u0438 \u0441\u043a\u043b\u0430\u0434\u0441\u043a\u043e\u0433\u043e \u0445\u043e\u0437\u044f\u0439\u0441\u0442\u0432\u0430", None))
        self.comboBox.setItemText(8, QCoreApplication.translate("MainWindow", u"I. \u0423\u0441\u043b\u0443\u0433\u0438 \u0433\u043e\u0441\u0442\u0438\u043d\u0438\u0447\u043d\u043e\u0433\u043e \u0445\u043e\u0437\u044f\u0439\u0441\u0442\u0432\u0430 \u0438 \u043e\u0431\u0449\u0435\u0441\u0442\u0432\u0435\u043d\u043d\u043e\u0433\u043e \u043f\u0438\u0442\u0430\u043d\u0438\u044f", None))
        self.comboBox.setItemText(9, QCoreApplication.translate("MainWindow", u"J. \u0423\u0441\u043b\u0443\u0433\u0438 \u0432 \u043e\u0431\u043b\u0430\u0441\u0442\u0438 \u0438\u043d\u0444\u043e\u0440\u043c\u0430\u0446\u0438\u0438 \u0438 \u0441\u0432\u044f\u0437\u0438", None))
        self.comboBox.setItemText(10, QCoreApplication.translate("MainWindow", u"K. \u0423\u0441\u043b\u0443\u0433\u0438 \u0444\u0438\u043d\u0430\u043d\u0441\u043e\u0432\u044b\u0435 \u0438 \u0441\u0442\u0440\u0430\u0445\u043e\u0432\u044b\u0435", None))
        self.comboBox.setItemText(11, QCoreApplication.translate("MainWindow", u"L. \u0423\u0441\u043b\u0443\u0433\u0438, \u0441\u0432\u044f\u0437\u0430\u043d\u043d\u044b\u0435 \u0441 \u043d\u0435\u0434\u0432\u0438\u0436\u0438\u043c\u044b\u043c \u0438\u043c\u0443\u0449\u0435\u0441\u0442\u0432\u043e\u043c", None))
        self.comboBox.setItemText(12, QCoreApplication.translate("MainWindow", u"M. \u0423\u0441\u043b\u0443\u0433\u0438, \u0441\u0432\u044f\u0437\u0430\u043d\u043d\u044b\u0435 \u0441 \u043d\u0430\u0443\u0447\u043d\u043e\u0439, \u0438\u043d\u0436\u0435\u043d\u0435\u0440\u043d\u043e-\u0442\u0435\u0445\u043d\u0438\u0447\u0435\u0441\u043a\u043e\u0439 \u0438 \u043f\u0440\u043e\u0444\u0435\u0441\u0441\u0438\u043e\u043d\u0430\u043b\u044c\u043d\u043e\u0439 \u0434\u0435\u044f\u0442\u0435\u043b\u044c\u043d\u043e\u0441\u0442\u044c\u044e", None))
        self.comboBox.setItemText(13, QCoreApplication.translate("MainWindow", u"N. \u0423\u0441\u043b\u0443\u0433\u0438 \u0430\u0434\u043c\u0438\u043d\u0438\u0441\u0442\u0440\u0430\u0442\u0438\u0432\u043d\u044b\u0435 \u0438 \u0432\u0441\u043f\u043e\u043c\u043e\u0433\u0430\u0442\u0435\u043b\u044c\u043d\u044b\u0435", None))
        self.comboBox.setItemText(14, QCoreApplication.translate("MainWindow", u"O. \u0423\u0441\u043b\u0443\u0433\u0438 \u0432 \u0441\u0444\u0435\u0440\u0435 \u0433\u043e\u0441\u0443\u0434\u0430\u0440\u0441\u0442\u0432\u0435\u043d\u043d\u043e\u0433\u043e \u0443\u043f\u0440\u0430\u0432\u043b\u0435\u043d\u0438\u044f \u0438 \u043e\u0431\u0435\u0441\u043f\u0435\u0447\u0435\u043d\u0438\u044f \u0432\u043e\u0435\u043d\u043d\u043e\u0439 \u0431\u0435\u0437\u043e\u043f\u0430\u0441\u043d\u043e\u0441\u0442\u0438; \u0443\u0441\u043b\u0443\u0433\u0438 \u043f\u043e \u043e\u0431\u044f\u0437\u0430\u0442\u0435\u043b\u044c\u043d\u043e\u043c\u0443 \u0441\u043e\u0446\u0438\u0430\u043b\u044c\u043d\u043e\u043c\u0443 \u043e\u0431\u0435\u0441\u043f\u0435\u0447\u0435\u043d\u0438\u044e", None))
        self.comboBox.setItemText(15, QCoreApplication.translate("MainWindow", u"P. \u0423\u0441\u043b\u0443\u0433\u0438 \u0432 \u043e\u0431\u043b\u0430\u0441\u0442\u0438 \u043e\u0431\u0440\u0430\u0437\u043e\u0432\u0430\u043d\u0438\u044f", None))
        self.comboBox.setItemText(16, QCoreApplication.translate("MainWindow", u"Q. \u0423\u0441\u043b\u0443\u0433\u0438 \u0432 \u043e\u0431\u043b\u0430\u0441\u0442\u0438 \u0437\u0434\u0440\u0430\u0432\u043e\u043e\u0445\u0440\u0430\u043d\u0435\u043d\u0438\u044f \u0438 \u0441\u043e\u0446\u0438\u0430\u043b\u044c\u043d\u044b\u0435 \u0443\u0441\u043b\u0443\u0433\u0438", None))
        self.comboBox.setItemText(17, QCoreApplication.translate("MainWindow", u"R. \u0423\u0441\u043b\u0443\u0433\u0438 \u0432 \u043e\u0431\u043b\u0430\u0441\u0442\u0438 \u0438\u0441\u043a\u0443\u0441\u0441\u0442\u0432\u0430, \u0440\u0430\u0437\u0432\u043b\u0435\u0447\u0435\u043d\u0438\u0439, \u043e\u0442\u0434\u044b\u0445\u0430 \u0438 \u0441\u043f\u043e\u0440\u0442\u0430", None))
        self.comboBox.setItemText(18, QCoreApplication.translate("MainWindow", u"S. \u041f\u0440\u043e\u0447\u0438\u0435 \u0443\u0441\u043b\u0443\u0433\u0438", None))
        self.comboBox.setItemText(19, QCoreApplication.translate("MainWindow", u"T. \u0422\u043e\u0432\u0430\u0440\u044b \u0438 \u0443\u0441\u043b\u0443\u0433\u0438 \u0440\u0430\u0437\u043b\u0438\u0447\u043d\u044b\u0435, \u043f\u0440\u043e\u0438\u0437\u0432\u043e\u0434\u0438\u043c\u044b\u0435 \u0434\u043e\u043c\u0430\u0448\u043d\u0438\u043c\u0438 \u0445\u043e\u0437\u044f\u0439\u0441\u0442\u0432\u0430\u043c\u0438 \u0434\u043b\u044f \u0441\u043e\u0431\u0441\u0442\u0432\u0435\u043d\u043d\u043e\u0433\u043e \u043f\u043e\u0442\u0440\u0435\u0431\u043b\u0435\u043d\u0438\u044f, \u0432\u043a\u043b\u044e\u0447\u0430\u044f \u0443\u0441\u043b\u0443\u0433\u0438 \u0440\u0430\u0431\u043e\u0442\u043e\u0434\u0430\u0442\u0435\u043b\u044f \u0434\u043b\u044f \u0434\u043e\u043c\u0430\u0448\u043d\u0435\u0433\u043e \u043f\u0435\u0440\u0441\u043e\u043d\u0430\u043b\u0430", None))
        self.comboBox.setItemText(20, QCoreApplication.translate("MainWindow", u"U. \u0423\u0441\u043b\u0443\u0433\u0438, \u043f\u0440\u0435\u0434\u043e\u0441\u0442\u0430\u0432\u043b\u044f\u0435\u043c\u044b\u0435 \u044d\u043a\u0441\u0442\u0435\u0440\u0440\u0438\u0442\u043e\u0440\u0438\u0430\u043b\u044c\u043d\u044b\u043c\u0438 \u043e\u0440\u0433\u0430\u043d\u0438\u0437\u0430\u0446\u0438\u044f\u043c\u0438 \u0438 \u043e\u0440\u0433\u0430\u043d\u0430\u043c\u0438", None))

hdr = {'user-agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) '
                     'Chrome/120.0.0.0 Safari/537.36'}


keywords_samples = ['компьютеры', 'компьютерная техника', 'расходные материалы', 'сервер',
                    'мультимедийная продукция',
                    'многофункциональное устройсво',
                    'системный блок',
                    'клавиатура',
                    'монитор',
                    'процессор'
                    ]

keywords_namecontract = ['поставка компьютеров',
                         'поставка АРМ',
                         'офисной техники',
                         'мультимедийного оборудования',
                         'серверов',
                         'клавиатур',
                         'системных блоков',
                         'монитор',
                         'запасные части',
                         'компьюетрного оборудования',
                         'компьютеры',
                         'суперкомпьютеры',
                         'сервера',
                         'принтеры',
                         'манипуляторы',
                         'компьютерные мыши']






procedure_number = []
customer = []
method_of_conducting = []
date_of_placement = []
end_date = []
nmc = []
electronic_platform = []
data_samples = []
data_analogs = []
data_delivery_time = []
data_payment_time = []
data_divisibility = []
data_address = []
data_support = []
data_url = []
data_dict = {'Номер процедуры': procedure_number,
             'Заказчик': customer,
             'Способ проведения': method_of_conducting,
             'Дата размещения': date_of_placement,
             'Дата окончания подачи заявок': end_date,
             'НМЦ': nmc,
             'Электронная площадка': electronic_platform,
             'Образцы': data_samples,
             'Аналоги': data_analogs,
             'Срок поставки': data_delivery_time,
             'Срок оплаты': data_payment_time,
             'Делимость': data_divisibility,
             'Адрес': data_address,
             'Обеспечение': data_support,
             'Ссылка': data_url}


def main():
    start_time = datetime.now()
    file1 = open("urls.txt", "r")
    while True:
        # считываем строку
        line = file1.readline()
        # прерываем цикл, если строка пустая
        if not line:
            break

        if os.path.isdir("content"):
            shutil.rmtree('content')
            os.mkdir("content")
        else:
            os.mkdir("content")
        print(line)
        if 'zakupki.gov.ru/epz/order/notice' in line:
            parse_epz_order(line)
        elif 'zakupki.gov.ru/223/purchase/public' in line:
            parse_223_order(line)
        elif 'zakupki.gov.ru/epz/pricereq/card' in line:
            parse_order_card(line)
        else:
            write_error_to_file(line.strip())

    # закрываем файл
    file1.close()
    write_dataframe_to_excel(data_dict)
    print(datetime.now() - start_time)


def parse_order_card(line):
    while True:
        try:
            procedure_str = ''
            customer_str = ''
            method_str = ''
            date_str = ''
            end_date_str = ''
            platform_str = ''
            nmc_str = ''
            url = line.strip()
            r = requests.get(url, headers=hdr)
            html_page = urlopen(url)
            soup = BeautifulSoup(html_page, features="lxml")
            # div_1 = soup.find("div", {"class": "span-body"})  # Parse the first div.
            # div_2 = div_1("div", {"class": "timestamp updated"})
            number = soup.find("span", {"class": "navBreadcrumb__text"})
            procedure_str += number.text.strip()
            # span = soup.find('span', {'class': "cardMainInfo__purchaseLink distancedText"})
            # print(span.text.strip())
            divs = soup.findAll('div', {'class': "cardMainInfo__section col-6"})
            for div in divs:
                spans = div.findAll('span')
                if spans[0].text.strip() == 'Размещено':
                    date_str += spans[1].text.strip()

            divs = soup.findAll('div', {'class': "cardMainInfo__section col-12"})
            for div in divs:
                spans = div.findAll('span')
                if spans[0].text.strip() == 'Сроки проведения закупки':
                    end_date_str += spans[0].text.strip() + ': ' + spans[1].text.strip()

            divs = soup.findAll('div', {'class': 'cardMainInfo__section'})
            for div in divs:
                spans = div.findAll('span')
                if spans[0].text.strip() == 'Организация, разместившая запрос цен':
                    customer_str += spans[0].text.strip() + ': ' + spans[1].text.strip()
            method_str += 'Процедура сбора информации'
            nmc_str += 'Не указано'
            platform_str += 'Не указано'
            add_parse_str_to_list(customer_str, date_str, end_date_str, method_str, nmc_str, platform_str,
                                  procedure_str)
            break
        except Exception:
            pass
    url_epz = line.replace('common-info.html', 'docs.html')
    while True:
        try:
            req = Request(
                url_epz,
                data=None,
                headers={
                    'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_9_3) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/35.0.1916.47 Safari/537.36'
                }
            )

            f = urlopen(req)
            html_page = f.read().decode('utf-8')
            soup = BeautifulSoup(html_page, features="lxml")
            for link_data in soup.findAll('div', {'class': 'price_file'}):
                link = link_data.findAll('a')[1]
                filename = link.text.strip()
                url_download = link.get('href')
                print(filename)
                while True:
                    try:
                        file_object = requests.get(url_download, headers=hdr)
                        filename = check_file_extension(file_object, filename)
                        with open(f'content/{filename}', 'wb') as local_file:
                            local_file.write(file_object.content)
                            local_file.close()
                            extract_files_from_archive(filename)
                        break
                    except Exception:
                        pass
            break
        except Exception as e:
            print(e)
            pass
    parse_docs_from_dir(url_epz)


def parse_223_order(line):
    doc_url, r, r_docs, url = parse_url(line.strip())
    tries = 0
    need_to_parse = True
    while True:
        try:
            html_page = get_html_page(doc_url)
            soup = BeautifulSoup(html_page, "lxml")
            need_to_parse = download_docs(soup, url)
            break
        except Exception:
            if tries >= 30:
                write_error_to_file(url)
                break
            tries += 1
            pass
    if need_to_parse:
        parse_docs_from_dir(url)


def parse_epz_order(line):
    ex = parse_url_epz(line)
    if ex:
        url_epz = line.replace('common-info.html', 'documents.html')
        print(url_epz)
        download_docs_epz(url_epz)
        parse_docs_from_dir(url_epz)


def download_docs_epz(url_epz):
    while True:
        try:
            req = Request(
                url_epz,
                data=None,
                headers={
                    'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_9_3) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/35.0.1916.47 Safari/537.36'
                }
            )

            f = urlopen(req)
            html_page = f.read().decode('utf-8')
            soup = BeautifulSoup(html_page, features="lxml")
            for link_data in soup.findAll('div', {'class': 'col clipText'}):
                link = link_data.findAll('a')[1]
                filename = link.text.strip()
                url_download = link.get('href')
                print(filename)
                while True:
                    try:
                        file_object = requests.get(url_download, headers=hdr)
                        filename = check_file_extension(file_object, filename)
                        with open(f'content/{filename}', 'wb') as local_file:
                            local_file.write(file_object.content)
                            local_file.close()
                            extract_files_from_archive(filename)
                        break
                    except Exception:
                        pass
            break
        except Exception as e:
            if e.args[0] == 'list index out of range':
                break
            pass


def parse_docs_from_dir(url):
    check_dirs = check_dirs_in_content()
    while check_dirs:
        check_dirs = check_dirs_in_content()
    # fix fix fix       doc to docx
    word = wc.Dispatch('Word.Application.8')
    folder = os.getcwd() + '\\content'
    for dir_path, dirs, files in os.walk(folder):
        for file_name in files:
            try:
                file_path = os.path.join(dir_path, file_name)
                file_name, file_extension = os.path.splitext(file_path)

                if "~$" not in file_name:
                    if file_extension.lower() == '.doc':  #
                        docx_file = '{0}{1}'.format(file_path, 'x')

                        if not os.path.isfile(docx_file):  # Skip conversion where docx file already exists

                            file_path = os.path.abspath(file_path)
                            docx_file = os.path.abspath(docx_file)
                            try:
                                word_doc = word.Documents.Open(file_path)
                                word_doc.SaveAs2(docx_file, FileFormat=16)
                                word_doc.Close()
                            except Exception as e:
                                print('Failed to Convert: {0}'.format(file_path))
                    if file_extension.lower() == '.pdf':
                        try:
                            cv = Converter(file_name + '.pdf')
                            cv.convert(file_name + '.docx', start=0, end=None)
                            cv.close()
                        except Exception:
                            pass
            except Exception:
                pass
    word.Quit()
    parse_docs(url)
    shutil.rmtree('content')


def parse_url_epz(line):
    while True:
        try:
            procedure_str = ''
            customer_str = ''
            method_str = ''
            date_str = ''
            end_date_str = ''
            platform_str = ''
            nmc_str = ''
            url = line.strip()
            r = requests.get(url, headers=hdr)
            html_page = urlopen(url)
            soup = BeautifulSoup(html_page, features="lxml")
            # div_1 = soup.find("div", {"class": "span-body"})  # Parse the first div.
            # div_2 = div_1("div", {"class": "timestamp updated"})
            number = soup.find("span", {"class": "navBreadcrumb__text"})
            procedure_str += number.text.strip()
            print(procedure_str)
            divs = soup.findAll('div', {"class": "cardMainInfo__section"})
            if len(divs) == 0:
                write_error_to_file(line.strip())
                return False
            for div in divs:
                spans = div.findAll('span')
                if spans[0].text.strip() == 'Заказчик' or spans[
                    0].text.strip() == 'Организация, осуществляющая размещение':
                    customer_str += spans[1].text.strip()
                if spans[0].text.strip() == 'Размещено':
                    date_str += spans[1].text.strip()
                if spans[0].text.strip() == 'Окончание подачи заявок':
                    end_date_str += spans[1].text.strip()
            divs = soup.find('div', {"class": "price"})
            spans = divs.findAll('span')
            if spans[0].text.strip() == 'Начальная цена':
                nmc_str += spans[1].text.strip()
            sections = soup.findAll('section', {'class': "blockInfo__section section"})
            for section in sections:
                spans = section.findAll('span')
                if 'Адрес электронной площадки в информационно-телекоммуникационной сети' in spans[0].text.strip():
                    platform_str += spans[1].text.strip()
            method_str += 'нет'
            add_parse_str_to_list(customer_str, date_str, end_date_str, method_str, nmc_str, platform_str,
                                  procedure_str)
            return True
        except IndexError:
            write_error_to_file(line.strip())
            return False
        except Exception as e:
            pass


def write_error_to_file(url):
    my_file = open("logs.txt", "a+")
    logs = f'{url}\t is not correct\t' + str(datetime.now()) + '\n'
    my_file.write(logs)
    my_file.close()


def write_dataframe_to_excel(data_to_excel):
    df = pd.DataFrame(data_to_excel)
    if os.path.isfile('data.xlsx'):
        old_df = pd.read_excel('data.xlsx', index_col=0)
        df = old_df.append(df, ignore_index=True)

    df.to_excel("data.xlsx")


def parse_docs(url):
    str_samples = ''
    str_analogs = ''
    str_delivery_time = ''
    str_payment_time = ''
    str_divisibility = ''
    str_address = ''
    str_support = ''

    for filename in os.listdir("content"):
        print(filename)
        split_tup = os.path.splitext(filename)
        extension = split_tup[1].lower()
        try:
            str_address, str_analogs, str_delivery_time, str_divisibility, str_payment_time, str_samples, str_support \
                = get_info_from_docx(extension, filename, str_address, str_analogs, str_delivery_time, str_divisibility,
                                     str_payment_time, str_samples, str_support)
        except Exception:
            pass

    add_data_to_list(str_address, str_analogs, str_delivery_time, str_divisibility, str_payment_time, str_samples,
                     str_support, url)


# def convert_pdf2docx(input_file: str, output_file: str):
#     """Преобразует PDF в DOCX"""
#     result = parse(pdf_file=input_file,
#                    docx_with_path=output_file)
#     return result


def get_info_from_docx(extension, filename, str_address, str_analogs, str_delivery_time, str_divisibility,
                       str_payment_time, str_samples, str_support):
    # if '.pdf' == extension:
    #     pdfFile = wi(filename=f'content/{filename}', resolution=300)
    #     image = pdfFile.convert('jpeg')
    #
    #     imageBlobs = []
    #
    #     for img in image.sequence:
    #         imgPage = wi(image=img)
    #         imageBlobs.append(imgPage.make_blob('jpeg'))
    #
    #     extract = []
    #
    #     for imgBlob in imageBlobs:
    #         image = Image.open(io.BytesIO(imgBlob))
    #         text = pytesseract.image_to_string(image, lang='eng')
    #         extract.append(text)
    #
    #     print(extract)
    # pdf = open(f'content/{filename}', 'rb')
    # pdf_reader = PyPDF2.PdfFileReader(pdf)
    # total_pages = pdf_reader.getNumPages()
    # text = ''
    # for i in range(total_pages):
    #     page = pdf_reader.getPage(0)
    #     text += page.extractText()

    # with pdfplumber.open(f'content/{filename}') as pdf:
    # total_pages = len(pdf.pages)
    # text = ''
    # for i in range(total_pages):
    #     page_obj = pdf.pages[i]
    #     text += page_obj.extract_text()

    # raw = parser.from_file(f'content/{filename}')
    # print(raw['content'])

    # split_tup = os.path.splitext(filename)
    # extension = '.docx'
    # new_filename = split_tup[0] + extension
    # convert_pdf2docx(f'content/{filename}', f'content/{new_filename}')

    # doc = aw.Document(f'content/{filename}')
    # print(filename + '\n\n')
    # split_tup = os.path.splitext(filename)
    # extension = '.docx'
    # filename = split_tup[0] + extension
    # print(filename)
    # doc.save(f'content/{filename}')
    if '.docx' == extension:
        try:
            doc = Document(f'content/{filename}')
        except Exception:
            return
        check_head = []
        for para in doc.paragraphs:
            try:
                str_address, str_analogs, str_delivery_time, str_divisibility, str_payment_time, str_samples, str_support = find_keywords(
                    para, str_address, str_analogs, str_delivery_time, str_divisibility, str_payment_time, str_samples,
                    str_support, filename)
                if "Heading" not in para.style.name:
                    run_bold_text = ''
                    for run in para.runs:
                        if run.bold:
                            run_bold_text = run_bold_text + ' ' + run.text
                            run_bold_text = run_bold_text.strip()
                    # print(run_bold_text)
                    if run_bold_text != '':
                        check_head = check_content_in_headers(run_bold_text)
                        # print(check_head)
                    if check_head and check_head[0]:
                        if check_head[1] == 'samples':
                            if filename not in str_samples:
                                str_samples += filename.strip() + ':\n'
                            str_samples += '\n' + para.text.strip()
                        if check_head[1] == 'analogs':
                            if filename not in str_analogs:
                                str_analogs += filename.strip() + ':\n'
                            str_analogs += '\n' + para.text.strip()
                        if check_head[1] == 'delivery_time':
                            if filename not in str_delivery_time:
                                str_delivery_time += filename.strip() + ':\n'
                            str_delivery_time += '\n' + para.text.strip()
                        if check_head[1] == 'payment_time':
                            if filename not in str_payment_time:
                                str_payment_time += filename.strip() + ':\n'
                            str_payment_time += '\n' + para.text.strip()
                        if check_head[1] == 'address':
                            if filename not in str_address:
                                str_address += filename.strip() + ':\n'
                            str_address += '\n' + para.text.strip()
                        if check_head[1] == 'divisibility':
                            if filename not in str_divisibility:
                                str_divisibility += filename.strip() + ':\n'
                            str_divisibility += '\n' + para.text.strip()
                        if check_head[1] == 'support':
                            if filename not in str_support:
                                str_support += filename.strip() + ':\n'
                            str_support += '\n' + para.text.strip()
            except Exception as e:
                print(e)
        for table in doc.tables:
            try:
                for row in table.rows:
                    for cell in row.cells:
                        for key_word in keywords_samples:
                            str_samples = check_length_text(cell, key_word, row, str_samples, filename)
                        for key_word in keywords_analogs:
                            str_analogs = check_length_text(cell, key_word, row, str_analogs, filename)
                        for key_word in keywords_delivery_time:
                            str_delivery_time = check_length_text(cell, key_word, row, str_delivery_time, filename)
                        for key_word in keywords_payment_time:
                            str_payment_time = check_length_text(cell, key_word, row, str_payment_time, filename)
                        for key_word in keywords_address:
                            str_address = check_length_text(cell, key_word, row, str_address, filename)
                        for key_word in keywords_divisibility:
                            str_divisibility = check_length_text(cell, key_word, row, str_divisibility, filename)
                        for key_word in keywords_support:
                            str_support = check_length_text(cell, key_word, row, str_support, filename)

                # Data will be a list of rows represented as dictionaries
                # containing each row's data.
                keys = None
                for i, row in enumerate(table.rows):
                    text = (cell.text.replace('\n', ' ') for cell in row.cells)
                    # Establish the mapping based on the first row
                    # headers; these will become the keys of our dictionary
                    if i == 0:
                        keys = tuple(text)
                        continue
                    # Construct a dictionary for this row, mapping
                    # keys to values for this row
                    row_data = dict(zip(keys, text))
                    for data in row_data:
                        str_delivery_time = get_vertical_info_from_table(data, row_data, str_delivery_time,
                                                                         keywords_delivery_time, filename)
                        str_samples = get_vertical_info_from_table(data, row_data, str_samples,
                                                                   keywords_samples, filename)
                        str_analogs = get_vertical_info_from_table(data, row_data, str_analogs,
                                                                   keywords_analogs, filename)
                        str_payment_time = get_vertical_info_from_table(data, row_data, str_payment_time,
                                                                        keywords_payment_time, filename)
                        str_divisibility = get_vertical_info_from_table(data, row_data, str_divisibility,
                                                                        keywords_divisibility, filename)
                        str_support = get_vertical_info_from_table(data, row_data, str_support,
                                                                   keywords_support, filename)
                        str_address = get_vertical_info_from_table(data, row_data, str_address,
                                                                   keywords_address, filename)
            except Exception as e:
                print(e)
    return str_address, str_analogs, str_delivery_time, str_divisibility, str_payment_time, str_samples, str_support


def get_vertical_info_from_table(data, row_data, str_param, keywords, filename):
    for key_word in keywords:
        if key_word in data:
            if data + ': ' + row_data[data] not in str_param:
                if filename not in str_param:
                    str_param += filename.strip() + ':\n'
                str_param += data + ': ' + row_data[data] + '\n'
    return str_param


def check_length_text(cell, key_word, row, str_param, filename):
    temp = ''
    if key_word in cell.text.lower():
        for cell_temp in row.cells:
            temp += '\n' + cell_temp.text.strip()
    if 300 > len(temp) > 0:
        if temp not in str_param:
            if filename not in str_param:
                str_param += filename.strip() + ':\n'
            str_param += '\n' + temp + '\n'
    return str_param


def check_content_in_headers(run_bold_text):
    res = []
    # образцы
    for key_word in keywords_samples:
        if key_word in run_bold_text.lower():
            res.append(True)
            res.append('samples')
            return res
    # аналоги
    for key_word in keywords_analogs:
        if key_word in run_bold_text.lower():
            res.append(True)
            res.append('analogs')
            return res
    # срок доставки
    for key_word in keywords_delivery_time:
        if key_word in run_bold_text.lower():
            res.append(True)
            res.append('delivery_time')
            return res
    # срок оплаты
    for key_word in keywords_payment_time:
        if key_word in run_bold_text.lower():
            res.append(True)
            res.append('payment_time')
            return res
    # адрес
    for key_word in keywords_address:
        if key_word in run_bold_text.lower():
            res.append(True)
            res.append('address')
            return res
    # делимость
    for key_word in keywords_divisibility:
        if key_word in run_bold_text.lower():
            res.append(True)
            res.append('divisibility')
            return res
    # обеспечение
    for key_word in keywords_support:
        if key_word in run_bold_text.lower():
            res.append(True)
            res.append('support')
            return res
    if not res:
        res.append(False)
        res.append('none')
    return res


def find_keywords(para, str_address, str_analogs, str_delivery_time, str_divisibility, str_payment_time, str_samples,
                  str_support, filename):
    try:
        if para.text or not para.text.isspace():
            # образцы
            str_samples = get_paragraphs(para, str_samples, keywords_samples, filename)
            # аналоги
            str_analogs = get_paragraphs(para, str_analogs, keywords_analogs, filename)
            # срок доставки
            str_delivery_time = get_paragraphs(para, str_delivery_time, keywords_delivery_time, filename)
            # срок оплаты
            str_payment_time = get_paragraphs(para, str_payment_time, keywords_payment_time, filename)
            # адрес
            str_address = get_paragraphs(para, str_address, keywords_address, filename)
            # делимость
            str_divisibility = get_paragraphs(para, str_divisibility, keywords_divisibility, filename)
            # обеспечение
            str_support = get_paragraphs(para, str_support, keywords_support, filename)
    except Exception as e:
        print(e)
    return str_address, str_analogs, str_delivery_time, str_divisibility, str_payment_time, str_samples, str_support


def get_paragraphs(para, param_str, keywords_list, filename):
    for key_word in keywords_list:
        if key_word in para.text.lower():
            if filename not in param_str:
                param_str += filename.strip() + ':\n'
            if para.text + '\n' not in param_str:
                param_str += para.text + '\n'
    return param_str


def add_data_to_list(str_address, str_analogs, str_delivery_time, str_divisibility, str_payment_time, str_samples,
                     str_support, url):
    if str_samples == '':
        data_samples.append('нет')
    else:
        str_samples = re.sub('[\n]+', '\n', str_samples)
        data_samples.append(str_samples)
    if str_analogs == '':
        data_analogs.append('нет')
    else:
        str_analogs = re.sub('[\n]+', '\n', str_analogs)
        data_analogs.append(str_analogs)
    if str_delivery_time == '':
        data_delivery_time.append('нет')
    else:
        str_delivery_time = re.sub('[\n]+', '\n', str_delivery_time)
        data_delivery_time.append(str_delivery_time)
    if str_payment_time == '':
        data_payment_time.append('нет')
    else:
        str_payment_time = re.sub('[\n]+', '\n', str_payment_time)
        data_payment_time.append(str_payment_time)
    if str_address == '':
        data_address.append('нет')
    else:
        str_address = re.sub('[\n]+', '\n', str_address)
        data_address.append(str_address)
    if str_divisibility == '':
        data_divisibility.append('нет')
    else:
        str_divisibility = re.sub('[\n]+', '\n', str_divisibility)
        data_divisibility.append(str_divisibility)
    if str_support == '':
        data_support.append('нет')
    else:
        str_support = re.sub('[\n]+', '\n', str_support)
        data_support.append(str_support)

    data_url.append(url)


def check_dirs_in_content():
    check_dirs = False
    list_dir = os.listdir("content")
    for dir_temp in list_dir:
        path = 'content/'
        path += dir_temp
        if os.path.isdir(path):
            check_dirs = True
            allfiles = os.listdir(path)
            for f in allfiles:
                os.rename(path + '/' + f, 'content/' + f)
            os.removedirs(path)
        if '.zip' in dir_temp or '.rar' in dir_temp or '.7z' in dir_temp:
            extract_files_from_archive(dir_temp)
            check_dirs = True
    return check_dirs


def download_docs(soup, url):
    need_to_parse = True
    allowed_downloads, count_downloads = check_validity_docs(soup)
    if allowed_downloads == 0:
        need_to_parse = False
        write_error_to_file(url)
    for link in soup.findAll('a', {'class': 'epz_aware'}):
        filename = link.text
        filename = filename.replace('\n', '')
        while True:
            try:
                url_download = "https://zakupki.gov.ru" + link.get('href')
                file_object = requests.get(url_download, headers=hdr)
                filename = check_file_extension(file_object, filename)
                with open(f'content/{filename}', 'wb') as local_file:
                    local_file.write(file_object.content)
                    local_file.close()
                    extract_files_from_archive(filename)
                break
            except Exception as e:
                print(e)
        count_downloads += 1
        if count_downloads >= allowed_downloads:
            break
    return need_to_parse


def check_file_extension(file_object, filename):
    content_type = file_object.headers['Content-Disposition']
    filename_redirect = re.findall("filename=(.+)", content_type)[0]
    split_tup = os.path.splitext(filename_redirect.replace('"', ''))
    extension = split_tup[1]
    if extension not in filename:
        filename += extension
    return filename


def check_validity_docs(soup):
    count_downloads = 0
    allowed_downloads = 0
    for numb in soup.findAll('td', {'style': 'width: 15%'}):
        if "(недействующая)" in numb.text:
            break
        allowed_downloads += 1
    return allowed_downloads, count_downloads


def extract_files_from_archive(filename):
    try:
        if ".rar" in filename or ".zip" in filename:
            Archive(f'content/{filename}').extractall('content')
            os.remove(f'content/{filename}')
        if ".7z" in filename:
            shutil.register_unpack_format('7zip', ['.7z'], unpack_7zarchive)
            shutil.unpack_archive(f'content/{filename}', 'content')
            os.remove(f'content/{filename}')
    except Exception:
        os.remove(f'content/{filename}')


def get_html_page(doc_url):
    count_connections = 0
    while True:
        try:
            count_connections += 1
            if count_connections > 50:
                break
            html_page = urlopen(doc_url)
            return html_page
        except Exception:
            pass


def parse_url(curr_url):
    while True:
        try:
            procedure_str = ''
            customer_str = ''
            method_str = ''
            date_str = ''
            end_date_str = ''
            platform_str = ''
            nmc_str = ''

            url = curr_url
            r = requests.get(url, headers=hdr)
            html_page = urlopen(url)
            soup = BeautifulSoup(html_page, features="lxml")
            rows = soup.findAll('tr')
            for row in rows:
                if row.find('span'):
                    list_text = row.text.strip().split('\n')
                    res_text = []
                    for text in list_text:
                        if not text.isspace() and text:
                            res_text.append(text.strip())
                    if 'Реестровый номер извещения' in res_text[0]:
                        procedure_str += res_text[1] + '\n'
                    if 'Наименование организации' in res_text[0]:
                        customer_str += res_text[1] + '\n'
                    if 'Способ размещения закупки' in res_text[0]:
                        method_str += res_text[1] + '\n'
                    if 'Дата размещения извещения' in res_text[0]:
                        date_str += res_text[1] + '\n'
                    if 'Дата и время окончания подачи заявок' in res_text[0]:
                        end_date_str += res_text[2] + '\n'
                    if 'Адрес электронной площадки' in res_text[0]:
                        platform_str += res_text[1] + '\n'

                if row.find('td'):
                    list_text = row.text.strip().split('\n')
                    res_text = []
                    for text in list_text:
                        if not text.isspace() and text:
                            res_text.append(text.strip())
                    if 'Наименование организации' in row.text:
                        customer_str += res_text[1] + '\n'
                    if 'Дата размещения извещения' in row.text:
                        date_str += res_text[1] + '\n'
                    if 'Адрес электронной площадки' in row.text:
                        platform_str += res_text[1] + '\n'
            break
        except Exception:
            pass
    while True:
        try:
            nmc_url = url.replace("common-info.html", "lot-list.html")
            r = requests.get(nmc_url, headers=hdr)
            html_page = get_html_page(nmc_url)
            soup = BeautifulSoup(html_page, features="lxml")
            rows = soup.findAll('td')
            for row in rows:
                if 'Российский рубль' in row.text:
                    nmc_str += row.text.strip() + '\n'
            break
        except Exception:
            pass

    add_parse_str_to_list(customer_str, date_str, end_date_str, method_str, nmc_str, platform_str, procedure_str)
    doc_url = url.replace("common-info.html", "documents.html")
    r_docs = requests.get(doc_url, headers=hdr)
    return doc_url, r, r_docs, url


def add_parse_str_to_list(customer_str, date_str, end_date_str, method_str, nmc_str, platform_str, procedure_str):
    if not procedure_str:
        procedure_number.append('нет')
    else:
        procedure_number.append(procedure_str)
    if not customer_str:
        customer.append('нет')
    else:
        customer.append(customer_str)
    if not method_str:
        method_of_conducting.append('нет')
    else:
        method_of_conducting.append(method_str)
    if not date_str:
        date_of_placement.append('нет')
    else:
        date_of_placement.append(date_str)
    if not end_date_str:
        end_date.append('нет')
    else:
        end_date.append(end_date_str)
    if not platform_str:
        electronic_platform.append('нет')
    else:
        electronic_platform.append(platform_str)
    if not nmc_str:
        nmc.append('нет')
    else:
        nmc.append(nmc_str)


# Press the green button in the gutter to run the script.
if __name__ == '__main__':
    main()
